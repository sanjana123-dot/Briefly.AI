import io
import re
from typing import Dict, Any
# Heavy libraries (fitz, docx) are imported lazily inside functions for fast startup

def extract_text_from_pdf(file_bytes: bytes) -> str:
    """
    Extracts plain text from a PDF file using PyMuPDF (fitz).
    Handles multiple pages and extracts text block by block.
    """
    text_content = []
    try:
        import fitz  # PyMuPDF - lazy import for fast startup
        # Open PDF document from bytes stream
        doc = fitz.open(stream=file_bytes, filetype="pdf")
        for page_num in range(len(doc)):
            page = doc.load_page(page_num)
            page_text = page.get_text("text")
            if page_text.strip():
                text_content.append(page_text)
        doc.close()
    except Exception as e:
        raise ValueError(f"Error extracting text from PDF: {str(e)}")
    
    return "\n\n".join(text_content)

def extract_text_from_docx(file_bytes: bytes) -> str:
    """
    Extracts text from a Word Document (.docx) using python-docx.
    Iterates through paragraphs and tables to reconstruct text.
    """
    text_content = []
    try:
        import docx  # python-docx - lazy import for fast startup
        doc = docx.Document(io.BytesIO(file_bytes))
        # Extract from paragraphs
        for paragraph in doc.paragraphs:
            if paragraph.text.strip():
                text_content.append(paragraph.text)
        
        # Extract text from tables
        for table in doc.tables:
            for row in table.rows:
                row_text = []
                for cell in row.cells:
                    cell_text = cell.text.strip()
                    if cell_text:
                        row_text.append(cell_text)
                if row_text:
                    text_content.append(" | ".join(row_text))
                    
    except Exception as e:
        raise ValueError(f"Error extracting text from DOCX: {str(e)}")
        
    return "\n\n".join(text_content)

def extract_text_from_txt(file_bytes: bytes) -> str:
    """
    Decodes plain text from a TXT bytes object.
    Attempts UTF-8 and falls back to ISO-8859-1 (Latin-1) or CP1252 if decoding fails.
    """
    encodings = ["utf-8", "latin-1", "cp1252"]
    for encoding in encodings:
        try:
            return file_bytes.decode(encoding)
        except UnicodeDecodeError:
            continue
    raise ValueError("Failed to decode TXT file. Supported encodings: UTF-8, Latin-1, CP1252.")

def clean_text(text: str) -> str:
    """
    Cleans raw text extracted from documents:
    - Normalizes double/triple spacing and tabs
    - Keeps paragraph structures
    - Eliminates redundant blank lines
    - Strips leading/trailing space on each line
    - Fixes mid-word hyphen breaks caused by PDF page ends
    """
    if not text:
        return ""
    
    # Remove excessive horizontal spaces and tabs
    text = re.sub(r'[ \t]+', ' ', text)
    
    # Fix broken hyphens (e.g., "de- \n velopment" -> "development")
    text = re.sub(r'(\w+)-\s*\n\s*(\w+)', r'\1\2', text)
    
    # Normalize newline gaps: ensure paragraphs are separated by a single clean empty line
    # first, split lines, clean padding, and merge
    lines = [line.strip() for line in text.splitlines()]
    
    # Reconstruct document while avoiding consecutive runs of more than 2 empty lines
    cleaned_lines = []
    consecutive_empty = 0
    for line in lines:
        if not line:
            consecutive_empty += 1
            if consecutive_empty <= 1:
                cleaned_lines.append("")
        else:
            consecutive_empty = 0
            cleaned_lines.append(line)
            
    text = "\n".join(cleaned_lines)
    return text.strip()

def get_document_stats(text: str, summary_text: str = "") -> Dict[str, Any]:
    """
    Calculates key metrics of the text:
    - Word count
    - Character count
    - Estimated reading time (based on 200 WPM)
    - Compression details (if summary is provided)
    """
    word_count = len(text.split())
    char_count = len(text)
    
    # Average adult reading speed: 200 words per minute
    reading_time_mins = max(1.0, round(word_count / 200.0, 1))
    
    stats = {
        "word_count": word_count,
        "char_count": char_count,
        "reading_time_mins": reading_time_mins
    }
    
    if summary_text:
        summary_words = len(summary_text.split())
        summary_chars = len(summary_text)
        
        reduction = 0.0
        if word_count > 0:
            reduction = ((word_count - summary_words) / word_count) * 100
            
        compression_ratio = 0.0
        if summary_words > 0:
            compression_ratio = word_count / summary_words
            
        stats.update({
            "summary_word_count": summary_words,
            "summary_char_count": summary_chars,
            "reduction_percentage": round(reduction, 1),
            "compression_ratio": round(compression_ratio, 2)
        })
        
    return stats

# Helper function to generate PDF bytes safely
def get_pdf_download_bytes(title: str, summary: str, stats: dict, keywords: list) -> bytes:
    from fpdf import FPDF
    
    class SummaryPDF(FPDF):
        def header(self):
            self.set_font('helvetica', 'B', 10)
            self.set_text_color(100, 110, 130)
            self.cell(0, 10, 'AI-POWERED DOCUMENT SUMMARIZER REPORT', border=False, ln=True, align='L')
            self.set_draw_color(124, 58, 237)
            self.line(10, 18, 200, 18)
            self.ln(8)
            
        def footer(self):
            self.set_y(-15)
            self.set_font('helvetica', 'I', 8)
            self.set_text_color(150, 150, 150)
            self.cell(0, 10, f'Page {self.page_no()}/{{nb}} | Generated by Antigravity AI Platform', align='C')

    try:
        pdf = SummaryPDF()
        pdf.alias_nb_pages()
        pdf.add_page()
        
        # Document Title
        pdf.set_font('helvetica', 'B', 16)
        pdf.set_text_color(30, 27, 75)
        cleaned_title = title.encode('latin-1', 'replace').decode('latin-1')
        pdf.multi_cell(0, 10, f"Document: {cleaned_title}")
        pdf.ln(4)
        
        # Metadata / Stats Header
        pdf.set_font('helvetica', 'B', 11)
        pdf.set_text_color(79, 70, 229)
        pdf.cell(0, 8, "1. DOCUMENT ANALYTICS", ln=True)
        pdf.set_font('helvetica', '', 9)
        pdf.set_text_color(40, 40, 40)
        
        pdf.cell(0, 6, f"- Original Word Count: {stats.get('word_count', 0)}", ln=True)
        pdf.cell(0, 6, f"- Summary Word Count: {stats.get('summary_word_count', 0)}", ln=True)
        pdf.cell(0, 6, f"- Volume Reduction: {stats.get('reduction_percentage', 0)}%", ln=True)
        pdf.cell(0, 6, f"- Compression Ratio: {stats.get('compression_ratio', 0)}x", ln=True)
        pdf.cell(0, 6, f"- Estimated Reading Time: {stats.get('reading_time_mins', 0)} minutes", ln=True)
        pdf.ln(6)
        
        # Keywords
        if keywords:
            pdf.set_font('helvetica', 'B', 11)
            pdf.set_text_color(79, 70, 229)
            pdf.cell(0, 8, "2. KEY PHRASES & TOPICS", ln=True)
            pdf.set_font('helvetica', '', 9)
            pdf.set_text_color(40, 40, 40)
            kw_str = ", ".join([f"{kw[0]}" for kw in keywords])
            cleaned_kw = kw_str.encode('latin-1', 'replace').decode('latin-1')
            pdf.multi_cell(0, 6, cleaned_kw)
            pdf.ln(6)
            
        # Summary Content
        pdf.set_font('helvetica', 'B', 12)
        pdf.set_text_color(79, 70, 229)
        pdf.cell(0, 10, "3. EXECUTIVE SUMMARY SYNTHESIS", ln=True)
        pdf.set_font('helvetica', '', 10)
        pdf.set_text_color(0, 0, 0)
        
        # Prepare text for latin-1 safety, substitute indicators, strip complex emojis
        clean_summary = summary.replace("•", "-").replace("💡", "* ").replace("🔑", "* ").replace("📌", "* ").replace("⚡", "* ").replace("🔍", "* ").replace("🎯", "* ").replace("📈", "* ").replace("⚙️", "* ")
        # General emojis removal
        clean_summary = re.sub(r'[^\x00-\x7F]+', ' ', clean_summary)
        
        cleaned_summary_latin = clean_summary.encode('latin-1', 'replace').decode('latin-1')
        pdf.multi_cell(0, 6, cleaned_summary_latin)
        
        # Convert bytearray/output to bytes to avoid Streamlit unsupported_error
        return bytes(pdf.output())
    except Exception as e:
        print(f"Error creating PDF: {str(e)}")
        return b""

# Helper function to generate DOCX bytes safely
def get_docx_download_bytes(title: str, summary: str, stats: dict, keywords: list) -> bytes:
    try:
        import docx  # python-docx - lazy import for fast startup
        doc = docx.Document()
        doc.add_heading('AI Document Summarizer Report', level=0)
        
        doc.add_heading('Source Document Details', level=1)
        doc.add_paragraph(f"Title: {title}")
        
        # Statistics Table
        doc.add_heading('Document Analytics', level=2)
        table = doc.add_table(rows=1, cols=2)
        table.style = 'Light Shading Accent 1'
        hdr_cells = table.rows[0].cells
        hdr_cells[0].text = 'Analytics Metric'
        hdr_cells[1].text = 'Value'
        
        metrics = [
            ("Original Word Count", str(stats.get('word_count', 0))),
            ("Summary Word Count", str(stats.get('summary_word_count', 0))),
            ("Reduction Percentage", f"{stats.get('reduction_percentage', 0)}%"),
            ("Compression Ratio", f"{stats.get('compression_ratio', 0)}x"),
            ("Est. Original Reading Time", f"{stats.get('reading_time_mins', 0)} mins")
        ]
        for metric, value in metrics:
            row_cells = table.add_row().cells
            row_cells[0].text = metric
            row_cells[1].text = value
            
        # Keywords
        if keywords:
            doc.add_heading('Top Keywords & Phrases', level=2)
            kw_str = ", ".join([f"{kw[0]} ({round(kw[1], 2)})" for kw in keywords])
            doc.add_paragraph(kw_str)
            
        doc.add_heading('Summary Synthesis', level=1)
        # Split summary text by paragraphs
        paragraphs = summary.split("\n\n")
        for para in paragraphs:
            if para.strip():
                # Check for bullet marks
                if para.strip().startswith("•") or para.strip().startswith("-"):
                    bullets = para.split("\n")
                    for bullet in bullets:
                        cleaned_bullet = bullet.strip().lstrip("•-").strip()
                        if cleaned_bullet:
                            doc.add_paragraph(cleaned_bullet, style='List Bullet')
                else:
                    doc.add_paragraph(para.strip())
                    
        doc_io = io.BytesIO()
        doc.save(doc_io)
        # Ensure returned value is strictly of type bytes to avoid Streamlit unsupported_error
        return bytes(doc_io.getvalue())
    except Exception as e:
        print(f"Error creating DOCX: {str(e)}")
        return b""
